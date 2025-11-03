"""
Export Utilities
Export use cases to various formats: DOCX, PlantUML, Markdown
"""
import os
from typing import List, Dict, Optional
from datetime import datetime


def export_to_docx(use_cases: List[Dict], session_context: Optional[Dict], session_id: str) -> str:
    """
    Export use cases to Microsoft Word document
    
    Args:
        use_cases: List of use cases
        session_context: Session metadata
        session_id: Session identifier
        
    Returns:
        Path to generated file
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('Use Case Specification', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    if session_context:
        doc.add_paragraph(f"Project: {session_context.get('project_context', 'N/A')}")
        doc.add_paragraph(f"Domain: {session_context.get('domain', 'N/A')}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Total Use Cases: {len(use_cases)}")
    doc.add_paragraph()  # Blank line
    
    # Add table of contents
    doc.add_heading('Table of Contents', 1)
    for idx, uc in enumerate(use_cases, 1):
        doc.add_paragraph(f"{idx}. {uc['title']}", style='List Number')
    
    doc.add_page_break()
    
    # Add each use case
    for idx, uc in enumerate(use_cases, 1):
        # Use case title
        doc.add_heading(f"{idx}. {uc['title']}", 1)
        
        # Preconditions
        doc.add_heading('Preconditions', 2)
        if uc.get('preconditions'):
            for precond in uc['preconditions']:
                doc.add_paragraph(precond, style='List Bullet')
        else:
            doc.add_paragraph('None specified', style='List Bullet')
        
        # Main Flow
        doc.add_heading('Main Flow', 2)
        if uc.get('main_flow'):
            for i, step in enumerate(uc['main_flow'], 1):
                doc.add_paragraph(f"{i}. {step}", style='List Number')
        else:
            doc.add_paragraph('Not specified')
        
        # Sub Flows
        doc.add_heading('Sub Flows (Optional Paths)', 2)
        if uc.get('sub_flows') and uc['sub_flows'] != ['No subflows']:
            for sub in uc['sub_flows']:
                doc.add_paragraph(sub, style='List Bullet')
        else:
            doc.add_paragraph('None specified')
        
        # Alternate Flows
        doc.add_heading('Alternate Flows (Error Handling)', 2)
        if uc.get('alternate_flows') and uc['alternate_flows'] != ['No alternate flows']:
            for alt in uc['alternate_flows']:
                doc.add_paragraph(alt, style='List Bullet')
        else:
            doc.add_paragraph('None specified')
        
        # Outcomes
        doc.add_heading('Expected Outcomes', 2)
        if uc.get('outcomes'):
            for outcome in uc['outcomes']:
                doc.add_paragraph(outcome, style='List Bullet')
        else:
            doc.add_paragraph('Not specified')
        
        # Stakeholders
        doc.add_heading('Stakeholders', 2)
        if uc.get('stakeholders'):
            stakeholder_text = ', '.join(uc['stakeholders'])
            doc.add_paragraph(stakeholder_text)
        else:
            doc.add_paragraph('Not specified')
        
        # Add page break between use cases (except last one)
        if idx < len(use_cases):
            doc.add_page_break()
    
    # Save document
    export_dir = '/tmp'
    os.makedirs(export_dir, exist_ok=True)
    file_path = os.path.join(export_dir, f'use_cases_{session_id}.docx')
    doc.save(file_path)
    
    return file_path


def export_to_plantuml(use_cases: List[Dict]) -> str:
    """
    Export use cases as PlantUML diagram
    
    Args:
        use_cases: List of use cases
        
    Returns:
        PlantUML code as string
    """
    plantuml = "@startuml\n"
    plantuml += "left to right direction\n"
    plantuml += "skinparam packageStyle rectangle\n\n"
    
    # Extract all unique stakeholders (actors)
    actors = set()
    for uc in use_cases:
        if uc.get('stakeholders'):
            actors.update(uc['stakeholders'])
    
    # Add actors
    plantuml += "' Actors\n"
    actor_map = {}
    for actor in sorted(actors):
        # Create valid PlantUML identifier
        actor_id = actor.replace(' ', '_').replace('-', '_')
        actor_map[actor] = actor_id
        
        # Use different notation for system vs human actors
        if 'system' in actor.lower() or 'database' in actor.lower() or 'api' in actor.lower():
            plantuml += f"rectangle {actor_id} as \"{actor}\"\n"
        else:
            plantuml += f"actor {actor_id} as \"{actor}\"\n"
    
    plantuml += "\n"
    
    # Add use cases
    plantuml += "' Use Cases\n"
    uc_map = {}
    for idx, uc in enumerate(use_cases):
        uc_id = f"UC{idx + 1}"
        uc_map[uc['title']] = uc_id
        
        # Escape quotes in title
        title = uc['title'].replace('"', '\\"')
        plantuml += f"usecase {uc_id} as \"{title}\"\n"
    
    plantuml += "\n"
    
    # Connect actors to use cases
    plantuml += "' Relationships\n"
    for uc in use_cases:
        uc_id = uc_map[uc['title']]
        if uc.get('stakeholders'):
            for stakeholder in uc['stakeholders']:
                actor_id = actor_map.get(stakeholder)
                if actor_id:
                    plantuml += f"{actor_id} --> {uc_id}\n"
    
    # Add dependencies between use cases if they share flows
    plantuml += "\n' Use Case Dependencies (include/extend)\n"
    for i, uc1 in enumerate(use_cases):
        uc1_id = uc_map[uc1['title']]
        for j, uc2 in enumerate(use_cases):
            if i >= j:
                continue
            
            uc2_id = uc_map[uc2['title']]
            
            # Check if uc2 is mentioned in uc1's flows
            uc1_text = ' '.join(uc1.get('main_flow', []) + uc1.get('sub_flows', []))
            if uc2['title'].lower() in uc1_text.lower():
                plantuml += f"{uc1_id} ..> {uc2_id} : <<include>>\n"
    
    plantuml += "\n@enduml"
    
    return plantuml


def export_to_markdown(use_cases: List[Dict], session_context: Optional[Dict], session_id: str) -> str:
    """
    Export use cases to Markdown document
    
    Args:
        use_cases: List of use cases
        session_context: Session metadata
        session_id: Session identifier
        
    Returns:
        Path to generated file
    """
    md = "# Use Case Specification\n\n"
    
    # Add metadata
    if session_context:
        md += f"**Project:** {session_context.get('project_context', 'N/A')}  \n"
        md += f"**Domain:** {session_context.get('domain', 'N/A')}  \n"
    md += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  \n"
    md += f"**Total Use Cases:** {len(use_cases)}  \n\n"
    
    # Add table of contents
    md += "## Table of Contents\n\n"
    for idx, uc in enumerate(use_cases, 1):
        # Create anchor link
        anchor = uc['title'].lower().replace(' ', '-').replace('/', '-')
        md += f"{idx}. [{uc['title']}](#{anchor})\n"
    md += "\n---\n\n"
    
    # Add each use case
    for idx, uc in enumerate(use_cases, 1):
        md += f"## {idx}. {uc['title']}\n\n"
        
        # Preconditions
        md += "### Preconditions\n\n"
        if uc.get('preconditions') and uc['preconditions'] != ['No preconditions']:
            for precond in uc['preconditions']:
                md += f"- {precond}\n"
        else:
            md += "*None specified*\n"
        md += "\n"
        
        # Main Flow
        md += "### Main Flow\n\n"
        if uc.get('main_flow') and uc['main_flow'] != ['No main flow']:
            for i, step in enumerate(uc['main_flow'], 1):
                md += f"{i}. {step}\n"
        else:
            md += "*Not specified*\n"
        md += "\n"
        
        # Sub Flows
        md += "### Sub Flows (Optional Paths)\n\n"
        if uc.get('sub_flows') and uc['sub_flows'] != ['No subflows']:
            for sub in uc['sub_flows']:
                md += f"- {sub}\n"
        else:
            md += "*None specified*\n"
        md += "\n"
        
        # Alternate Flows
        md += "### Alternate Flows (Error Handling)\n\n"
        if uc.get('alternate_flows') and uc['alternate_flows'] != ['No alternate flows']:
            for alt in uc['alternate_flows']:
                md += f"- {alt}\n"
        else:
            md += "*None specified*\n"
        md += "\n"
        
        # Outcomes
        md += "### Expected Outcomes\n\n"
        if uc.get('outcomes') and uc['outcomes'] != ['No outcomes']:
            for outcome in uc['outcomes']:
                md += f"- {outcome}\n"
        else:
            md += "*Not specified*\n"
        md += "\n"
        
        # Stakeholders
        md += "### Stakeholders\n\n"
        if uc.get('stakeholders') and uc['stakeholders'] != ['No stakeholders']:
            stakeholder_list = ', '.join(uc['stakeholders'])
            md += f"{stakeholder_list}\n"
        else:
            md += "*Not specified*\n"
        md += "\n"
        
        # Separator between use cases
        if idx < len(use_cases):
            md += "---\n\n"
    
    # Save markdown file
    export_dir = '/tmp'
    os.makedirs(export_dir, exist_ok=True)
    file_path = os.path.join(export_dir, f'use_cases_{session_id}.md')
    
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(md)
    
    return file_path


def export_to_json(use_cases: List[Dict], session_context: Optional[Dict]) -> Dict:
    """
    Export use cases to structured JSON format
    
    Args:
        use_cases: List of use cases
        session_context: Session metadata
        
    Returns:
        Structured JSON object
    """
    return {
        "metadata": {
            "project_context": session_context.get('project_context', '') if session_context else '',
            "domain": session_context.get('domain', '') if session_context else '',
            "generated_at": datetime.now().isoformat(),
            "total_use_cases": len(use_cases)
        },
        "use_cases": use_cases
    }
